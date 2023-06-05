# pip3 install python-docx -i http://mirrors.aliyun.com/pypi/simple --trusted-host mirrors.aliyun.com
# 不用--trusted-host就会报错：WARNING: Retrying (Retry(total=4, connect=None, read=None, redirect=None, status=None)) after connection broken by 'SSLError(SSLEOFError(8, 'EOF occurred in violation of protocol (_ssl.c:1123)'))': /simple/python-docx/
# 注意：不支持dox文档

# 参考：https://github.com/datawhalechina/office-automation/blob/main/Task03-Python%E4%B8%8EWord%E5%92%8CPDF/python%E4%B8%8Eword.md

# 导入库
import os

try:
    from docx import Document
except:
    os.system('pip install python-docx -i http://mirrors.aliyun.com/pypi/simple --trusted-host mirrors.aliyun.com')
    from docx import Document


# 新建空白文档
doc_1 = Document()

# 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9）
doc_1.add_heading('新建空白文档标题，级别为0', level=0)  # 级别为0在导航栏中不会显示
doc_1.add_heading('新建空白文档标题，级别为1', level=1)
doc_1.add_heading('新建空白文档标题，级别为2', level=2)

# 新增段落
paragraph_1 = doc_1.add_paragraph('这是第一段文字的开始\n请多多关照！')
# 加粗；Run-文字块
paragraph_1.add_run('加粗字体').bold = True
paragraph_1.add_run('普通字体')
# 斜体
paragraph_1.add_run('斜体字体').italic = True  # 仿意大利古书写体的、斜体的

# 新段落（当前段落的下方）
paragraph_2 = doc_1.add_paragraph('新起的第二段文字。')

# 新段落（指定端的上方）：在第一段的前面插入段落
prior_paragraph = paragraph_1.insert_paragraph_before('在第一段文字前插入的段落')

# 添加分页符(可以进行灵活的排版）
doc_1.add_page_break()


# 新段落（指定端的上方）
paragraph_3 = doc_1.add_paragraph('这是第二页第一段文字！')

# 保存文件（当前目录下）
doc_1.save('doc_1.docx')


#%%
# 导入库
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# 新建文档（Datawhale）
doc_1 = Document()

# 字体设置（全局）
'''只更改font.name是不够的，还需要调用._element.rPr.rFonts的set()方法。'''
doc_1.styles['Normal'].font.name = u'宋体'
doc_1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 添加标题（0相当于文章的题目，默认级别是1，级别范围为0-9，0时候自动带下划线）
# WD_ALIGN_PARAGRAPH. LEFT：左对齐；
# WD_ALIGN_PARAGRAPH. CENTER：居中对其；
# WD_ALIGN_PARAGRAPH. RIGHT：右对齐；
# WD_ALIGN_PARAGRAPH. JUSTIFY：两端对齐；

heading_1 = doc_1.add_heading('周杰伦', level=0)
heading_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐，默认居左对齐

# 新增段落
paragraph_1 = doc_1.add_paragraph()
'''
设置段落格式：首行缩进0.75cm，居左，段后距离1.0英寸,1.5倍行距。
'''
paragraph_1.paragraph_format.first_line_indent = Cm(0.75)
paragraph_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
paragraph_1.paragraph_format.space_after = Inches(1.0)
paragraph_1.paragraph_format.line_spacing = 1.5

text = '中国台湾华语流行歌手、' \
       '音乐创作家、作曲家、作词人、' \
       '制作人、杰威尔音乐公司老板之一、导演。' \
       '近年涉足电影行业。周杰伦是2000年后亚洲流行乐坛最具革命性与指标' \
       '性的创作歌手，有“亚洲流行天王”之称。他突破原有亚洲音乐的主题、形' \
       '式，融合多元的音乐素材，创造出多变的歌曲风格，尤以融合中西式曲风的嘻哈' \
       '或节奏蓝调最为著名，可说是开创华语流行音乐“中国风”的先声。周杰伦的' \
       '出现打破了亚洲流行乐坛长年停滞不前的局面，为亚洲流行乐坛翻开了新的一页！'

# 这样的字体设置是全局变量
r_1 = paragraph_1.add_run(text)
r_1.font.size = Pt(10)  # 字号
r_1.font.bold = True  # 加粗
r_1.font.color.rgb = RGBColor(255, 0, 0)  # 颜色

print(len(paragraph_1.runs))  # 查看段落拥有的run对象数量
print(paragraph_1.runs[0].text)  # 查看对应run对象的文本等属性

# 保存文件（当前目录下）
doc_1.save('周杰伦.docx')

#%% 字体设置 1
'''字体设置1.py'''
#导入库
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

document = Document() # 新建docx文档

# 设置宋体字样式
style_font = document.styles.add_style('宋体', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '宋体'
document.styles['宋体']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 设置楷体字样式
style_font = document.styles.add_style('楷体', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '楷体'
document.styles['楷体']._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

# 设置华文中宋字样式
style_font = document.styles.add_style('华文中宋', WD_STYLE_TYPE.CHARACTER)
style_font.font.name = '华文中宋'
document.styles['华文中宋']._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')

paragraph1 = document.add_paragraph() # 添加段落
run = paragraph1.add_run(u'aBCDefg这是中文', style='宋体') # 设置宋体样式

font = run.font  #设置字体
font.name = 'Cambira'  # 设置西文字体
paragraph1.add_run(u'aBCDefg这是中文', style='楷体').font.name = 'Cambira'
paragraph1.add_run(u'aBCDefg这是中文', style='华文中宋').font.name = 'Cambira'

document.save('字体设置1.docx')

#%% 字体设置2
'''字体设置2.py'''
#导入库
from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

#定义字体设置函数
def font_setting(doc,text,font_cn):
       style_add = doc.styles.add_style(font_cn, WD_STYLE_TYPE.CHARACTER)
       style_add.font.name = font_cn
       doc.styles[font_cn]._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
       par = doc.add_paragraph()
       text = par.add_run(text, style=font_cn)

doc = Document()
a = '小朋友 你是否有很多问号'
b = '为什么 别人在那看漫画'
c = '我却在学画画 对着钢琴说话'

font_setting(doc,a,'宋体')
font_setting(doc,b,'华文中宋')
font_setting(doc,c,'黑体')

doc.save('字体设置2.docx')


#%% 3.1.3插入图片与表格
#导入库
from docx import Document
from docx.shared import Inches

#打开已有的文档，在原来的文档上增加内容
doc_1 = Document('周杰伦.docx')   #上面脚本存储的文档
#新增图片
doc_1.add_picture('周杰伦.jpg',width=Inches(1.0), height=Inches(1.0))

# 创建3行1列表格
table1 = doc_1.add_table(rows=2, cols=1)
table1.style = 'Medium Grid 1 Accent 1'  #表格样式很多种，如，Light Shading Accent 1等

# 修改第0行第0列单元格的内容为营口
table1.cell(0, 0).text = '营口'
# 修改第1行第0列单元格的内容为人民
table1.rows[1].cells[0].text = '人民'

# 在表格底部新增一行
row_cells = table1.add_row().cells
# 新增行的第一列添加内容
row_cells[0].text = '加油'

doc_1.save('周杰伦为营口加油.docx')


#%% 3.1.4设置页眉页脚
# 在python-docx包中则要使用节(section)中的页眉(header)和页脚(footer)对象来具体设置。

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document() # 新建文档

header = document.sections[0].header # 获取第一个节的页眉
print('页眉中默认段落数：', len(header.paragraphs))
paragraph = header.paragraphs[0] # 获取页眉的第一个段落
paragraph.add_run('这是第一节的页眉') # 添加页面内容
footer = document.sections[0].footer # 获取第一个节的页脚
paragraph = footer.paragraphs[0] # 获取页脚的第一个段落
paragraph.add_run('这是第一节的页脚') # 添加页脚内容


'''在docx文档中又添加了2个节，共计3个节，页面和页脚会显示了“与上一节相同”。
如果不使用上一节的内容和样式要将header.is_linked_to_previous的属性或footer.is_linked_to_previous的属性设置为False，
用于解除“链接上一节页眉”或者“链接上一节页脚”。'''
document.add_section() # 添加一个新的节
document.add_section() # 添加第3个节
header = document.sections[1].header # 获取第2个节的页眉
header.is_linked_to_previous = False # 不使用上节内容和样式

#对齐设置
header = document.sections[1].header # 获取第2个节的页眉
header.is_linked_to_previous = False # 不使用上节内容和样式
paragraph = header.paragraphs[0]
paragraph.add_run('这是第二节的页眉')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 设置页眉居中对齐
document.sections[1].footer.is_linked_to_previous = False
footer.paragraphs[0].add_run('这是第二节的页脚') # 添加第2节页脚内容
footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # 设置第2节页脚居中对齐
header = document.sections[2].header # 获取第3个节的页眉
header.is_linked_to_previous = False # 不使用上节的内容和样式
paragraph = header.paragraphs[0] # 获取页眉中的段落
paragraph.add_run('这是第三节的页眉')
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # 设置页眉右对齐
document.sections[2].footer.is_linked_to_previous = False
footer.paragraphs[0].add_run('这是第三节的页脚') # 添加第3节页脚内容
footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # 设置第3节页脚右对齐
document.save('页眉页脚1.docx') # 保存文档
